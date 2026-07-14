"""
generate_summary.py
-------------------
Generates a highly readable, spaced-out, 2-page executive summary.docx.
"""

import sys
from pathlib import Path
from datetime import date

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("python-docx is not installed. Run: pip install python-docx")
    sys.exit(1)


BASE_DIR = Path(__file__).parent
CHARTS_DIR = BASE_DIR / "charts"
OUTPUT_PATH = BASE_DIR / "summary.docx"
TODAY = date.today().strftime("%d %b %Y")


# --- Helpers ---
def set_cell_bg(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E4053")
    pBdr.append(bottom)
    p._p.get_or_add_pPr().append(pBdr)

def heading(doc, text: str, space_before: int = 14):
    h = doc.add_heading(text, level=1)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after = Pt(2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(30, 60, 114)
        run.font.size = Pt(15)
        run.bold = True
    return h

def body(doc, text: str, bold_prefix: str = "", space_after: int = 8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
        r.font.color.rgb = RGBColor(30, 60, 114)
    r = p.add_run(text)
    r.font.size = Pt(12)
    return p

def bullet(doc, text: str, bold_prefix: str = ""):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
    r = p.add_run(text)
    r.font.size = Pt(12)

def chart_in_cell(cell, filename: str, caption: str, width: float = 3.2):
    path = CHARTS_DIR / filename
    if not path.exists(): return
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = cell.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    for run in cap.runs:
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)


# --- Build Document ---
def build():
    doc = Document()

    # Slightly relaxed margins to allow content to breathe and fill 2 pages perfectly
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(12)

    # HEADER
    title = doc.add_heading("Sales Forecasting & Demand Intelligence", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    for run in title.runs:
        run.font.color.rgb = RGBColor(30, 60, 114)
        run.font.size = Pt(22)
        run.bold = True

    # 1. EXECUTIVE SUMMARY
    heading(doc, "1. Executive Summary", space_before=6)
    body(doc, 
        "This report analyzes four years of sales data across three product lines and four regions. "
        "Overall growth is steady, driven by the Technology category and the West region. We anticipate "
        "a reliable revenue surge in Nov/Dec, consistent with historical trends. However, past data reveals "
        "inventory strain during unexpected demand spikes. By implementing predictive stock adjustments and tracking "
        "anomalies early, the supply chain can meet Q4 demand while minimizing excess warehouse costs.",
        space_after=10
    )

    # 2. KEY FINDINGS
    heading(doc, "2. Key Findings from Sales History")
    bullet(doc, "Technology is the primary revenue driver, outpacing Furniture and Office Supplies.", bold_prefix="Technology Leads:")
    bullet(doc, "The West region shows the strongest year-over-year growth.", bold_prefix="West Region Growth:")
    bullet(doc, "Shipping times average 3.9 days but exceed 5 days in lagging regions, impacting retention.", bold_prefix="Fulfillment Bottlenecks:")
    
    # 3. SALES FORECAST
    heading(doc, "3. Three-Month Sales Forecast (Oct - Dec)")
    body(doc, 
        "Our predictive model projects strong growth over the next 90 days. Because real-world demand fluctuates, "
        "these forecasts come with a 'confidence range'—giving us a safe upper and lower bound for planning inventory.",
        space_after=8
    )
    
    ftbl = doc.add_table(rows=4, cols=3)
    ftbl.style = "Table Grid"
    ftbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Month", "Expected Revenue", "Confidence Range (Low to High)"]
    for i, h in enumerate(headers):
        c = ftbl.rows[0].cells[i]
        set_cell_bg(c, "1E3C72")
        r = c.paragraphs[0].add_run(h)
        r.bold = True; r.font.color.rgb = RGBColor(255, 255, 255); r.font.size = Pt(11)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    rows_data = [
        ["October", "~238,500", "190,000  to  275,000"],
        ["November", "~252,000", "205,000  to  295,000"],
        ["December", "~275,000", "225,000  to  320,000"],
    ]
    for ri, row in enumerate(rows_data):
        for ci, val in enumerate(row):
            c = ftbl.rows[ri + 1].cells[ci]
            set_cell_bg(c, ["FFFFFF", "F2F4F4", "FFFFFF"][ri])
            r = c.paragraphs[0].add_run(val)
            r.font.size = Pt(11)
            if ci == 0: r.bold = True
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Adding 2 side-by-side charts for Forecast
    tbl_charts = doc.add_table(rows=1, cols=2)
    tbl_charts.alignment = WD_TABLE_ALIGNMENT.CENTER
    chart_in_cell(tbl_charts.cell(0, 0), "monthly_sales_trend.png", "Fig 1: Clear Upward Sales Trend & Strong Q4 Seasonality", 3.52)
    chart_in_cell(tbl_charts.cell(0, 1), "prophet_forecast.png", "Fig 2: Q4 Revenue Forecast with Confidence Bounds", 3.52)
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 4. ANOMALIES
    heading(doc, "4. Top 3 Demand Anomalies Detected", space_before=10)
    body(doc, "We reviewed historical data to find unexpected demand shocks that strained inventory:")
    bullet(doc, "A massive 61% spike above average, likely a major promotional event. Strains supply chains if not pre-planned.", bold_prefix="November 2016 Surge:")
    bullet(doc, "A sharp 38% drop in expected orders, likely a system outage or unrecorded sales.", bold_prefix="March 2015 Revenue Drop:")
    bullet(doc, "A sudden wave of orders, matching a classic B2B stockpiling pattern before Q4.", bold_prefix="September 2017 Institutional Buy:")

    # Adding 2 side-by-side charts for Anomalies
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    tbl_anomalies = doc.add_table(rows=1, cols=2)
    tbl_anomalies.alignment = WD_TABLE_ALIGNMENT.CENTER
    chart_in_cell(tbl_anomalies.cell(0, 0), "anomaly_isolation_forest.png", "Fig 3: Unforeseen Order Spikes", 3.52)
    chart_in_cell(tbl_anomalies.cell(0, 1), "shipping_time_by_region.png", "Fig 4: Shipping Delays Impacting Fulfillment", 3.52)
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 5. SEGMENTATION
    heading(doc, "5. Product Demand Segmentation & Strategy")
    body(doc, "By analyzing product volume, order frequency, and volatility, we segmented inventory into four groups:")
    
    ctbl = doc.add_table(rows=5, cols=3)
    ctbl.style = "Table Grid"
    ctbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    chdrs = ["Product Group", "Behavior", "Recommended Stocking Strategy"]
    for i, h in enumerate(chdrs):
        c = ctbl.rows[0].cells[i]
        set_cell_bg(c, "2E4053")
        r = c.paragraphs[0].add_run(h)
        r.bold = True; r.font.color.rgb = RGBColor(255,255,255); r.font.size = Pt(11)
    
    cdata = [
        ["Stable High-Volume\n(Paper, Binders)", "Sells constantly. Low variance.", "Maintain deep safety stock. Automate re-ordering."],
        ["Low-Freq High-Value\n(Copiers, Machines)", "Sells rarely, but huge margins.", "Hold moderate stock. Negotiate faster supplier lead times."],
        ["Growing Demand\n(Accessories, Labels)", "Demand is steadily rising.", "Increase quarterly order sizes proactively."],
        ["Volatile / Declining\n(Tables, Supplies)", "Erratic demand, tying up cash.", "Shift to Just-In-Time ordering to free up warehouse space."],
    ]
    for ri, row in enumerate(cdata):
        for ci, val in enumerate(row):
            c = ctbl.rows[ri+1].cells[ci]
            set_cell_bg(c, ["D5F5E3", "D6EAF8", "FEF9E7", "FDEDEC"][ri])
            r = c.paragraphs[0].add_run(val)
            r.font.size = Pt(11)
            if ci == 0: r.bold = True
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 6. RECOMMENDATIONS
    heading(doc, "6. Concrete Business Recommendations", space_before=6)
    bullet(doc, "Nov/Dec generate 40-60% above average revenue. Increase Q4 Purchase Orders for Technology by 35% before Oct 1st to prevent stockouts.", bold_prefix="1. Pre-Stock Technology for Q4:")
    bullet(doc, "Deploy an automated weekly alert flagging the supply chain team if order velocity deviates 1.5x from expected, buying 5-7 days of reaction time.", bold_prefix="2. Implement Automated Shock Alerts:")
    bullet(doc, "Move volatile SKUs (like Tables) to a Just-In-Time fulfillment model, reserving premium warehouse space for high-volume runners.", bold_prefix="3. Shift Volatile Goods to Just-In-Time:")
    
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 7. LIMITATIONS
    heading(doc, "7. System Risk & Limitation")
    
    # Shaded risk box
    tbl_risk = doc.add_table(rows=1, cols=1)
    tbl_risk.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_risk = tbl_risk.cell(0, 0)
    set_cell_bg(cell_risk, "FEF9E7") # Light yellow
    p = cell_risk.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    run = p.add_run("Risk factor: Models rely entirely on historical data and cannot predict macroeconomic shocks or supply chain breakdowns. Combine forecasts with real-time operational context.")
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(120, 80, 0)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    


    doc.save(str(OUTPUT_PATH))
    print(f"Jargon-free, highly readable summary saved to: {OUTPUT_PATH}")

if __name__ == "__main__":
    build()
